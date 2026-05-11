#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中国党政机关公文格式自动化排版工具
依据 GB/T 9704-2012 标准
使用 python-docx 库生成规范格式的 Word 文档

模块结构：
  1. 页面设置
  2. 样式定义
  3. 文本解析
  4. 格式应用
  5. 图片处理
  6. 文件输出
"""

import re
import os
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Mm, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 1. 页面设置模块
# ============================================================

def setup_page(doc):
    """
    设置文档页面参数（GB/T 9704-2012）
    - 页边距：上37mm，下35mm，左28mm，右26mm
    - 纸张方向：纵向
    - 纸张大小：A4 (210mm x 297mm)
    """
    section = doc.sections[0]

    # 纸张方向：纵向
    section.orientation = WD_ORIENT.PORTRAIT

    # 纸张大小：A4
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # 页边距
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)

    # 页脚距离
    section.footer_distance = Cm(2.5)

    return section


# ============================================================
# 2. 样式定义模块
# ============================================================

# 字体名称常量
FONT_FANGSONG = '仿宋_GB2312'    # 正文
FONT_FANGSONG_ALT = '仿宋'       # 正文备选
FONT_HEITI = '黑体'              # 一级标题
FONT_KAITI = '楷体_GB2312'       # 二级标题
FONT_KAITI_ALT = '楷体'          # 二级标题备选
FONT_XIAOBIAOSONG = '方正小标宋简体'  # 主标题
FONT_SONGTI = '宋体'             # 页码

# 字号常量（磅值）
SIZE_ERHAO = Pt(22)      # 二号，用于主标题
SIZE_SANHAO = Pt(16)     # 三号，用于正文
SIZE_SIHAO = Pt(14)      # 四号，用于页码
SIZE_XIAOSI = Pt(12)     # 小四号，用于图题

# 行距
LINE_SPACING = Pt(28)    # 固定值28磅

# 页面有效宽度（A4宽减去左右边距：210 - 28 - 26 = 156mm）
PAGE_CONTENT_WIDTH_MM = 156
# 图片默认宽度：页面有效宽度的80%，约12.5cm
IMAGE_DEFAULT_WIDTH_CM = 12.5

# 公文类型定义
DOC_TYPES = {
    "通知": {"name": "通知", "greeting_no_indent": False},
    "决定": {"name": "决定", "greeting_no_indent": False},
    "意见": {"name": "意见", "greeting_no_indent": False},
    "报告": {"name": "报告", "greeting_no_indent": False},
    "请示": {"name": "请示", "greeting_no_indent": False},
    "批复": {"name": "批复", "greeting_no_indent": False},
    "函": {"name": "函", "greeting_no_indent": False},
    "讲话稿": {"name": "讲话稿", "greeting_no_indent": True},
    "会议纪要": {"name": "会议纪要", "greeting_no_indent": True},
}

# 问候语匹配正则：行首匹配问候词，且行末为冒号/逗号（称呼行特征）
RE_GREETING = re.compile(
    r'^(各位|同志们|朋友们|女士们|先生们|同志|朋友|各位领导|各位同事|各位来宾).+[：:，,]$'
)


def set_run_font(run, font_name, font_size, bold=False, color=None):
    """
    设置 run 的字体属性
    - font_name: 中文字体名称
    - font_size: 字号（Pt值）
    - bold: 是否加粗
    - color: 字体颜色（RGBColor），默认黑色
    """
    run.font.size = font_size
    run.font.bold = bold
    if color is None:
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run.font.color.rgb = color

    # 设置中文字体
    run.font.name = font_name
    # 通过 rPr 设置东亚字体（确保中文字体生效）
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_format(paragraph, alignment=None, first_line_indent=None,
                         line_spacing=LINE_SPACING, space_before=Pt(0),
                         space_after=Pt(0)):
    """
    设置段落格式
    - alignment: 对齐方式
    - first_line_indent: 首行缩进（如 2字符 用 Pt(32) 近似）
    - line_spacing: 行距
    - space_before: 段前距
    - space_after: 段后距
    """
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = space_before
    pf.space_after = space_after


def add_page_number(doc):
    """
    添加页码：页面底部居中，格式为 "- 1 -"
    字体：宋体，四号
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # 清空已有页脚内容
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # 新建页脚段落
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加页码字段 "- 页码 -"
    run1 = footer_para.add_run('- ')
    set_run_font(run1, FONT_SONGTI, SIZE_SIHAO)

    # 插入页码域代码
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_page = footer_para.add_run()
    run_page._element.append(fld_char_begin)
    set_run_font(run_page, FONT_SONGTI, SIZE_SIHAO)

    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr = footer_para.add_run()
    run_instr._element.append(instr_text)
    set_run_font(run_instr, FONT_SONGTI, SIZE_SIHAO)

    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = footer_para.add_run()
    run_end._element.append(fld_char_end)
    set_run_font(run_end, FONT_SONGTI, SIZE_SIHAO)

    run2 = footer_para.add_run(' -')
    set_run_font(run2, FONT_SONGTI, SIZE_SIHAO)


# ============================================================
# 3. 文本解析模块
# ============================================================

# 正则表达式：识别各层级标题
RE_LEVEL1 = re.compile(r'^[一二三四五六七八九十]+、')          # 一级标题：一、二、
RE_LEVEL2 = re.compile(r'^[（\(][一二三四五六七八九十]+[）\)]')  # 二级标题：（一）（二）
RE_LEVEL3 = re.compile(r'^\d+[\.．]')                          # 三级标题：1. 2.
RE_FAWEN_HAO = re.compile(r'^[^\s]*〔\d{4}〕\d+号')           # 发文字号
RE_TABLE_SEPARATOR = re.compile(r'^\|[-\s|]+\|$')              # 表格分隔行
RE_TABLE_ROW = re.compile(r'^\|.*\|$')                         # 表格数据行

# 图片标记正则：匹配 [图片1] [图片1,10cm] [图片1,50%] [图片1:图题] [图片1,10cm:图题]
# 也匹配 【图片1】 [图1] 【图1】等变体
RE_IMAGE = re.compile(
    r'[【\[](图片?\d+)[,\s]*(宽?(?:\d+(?:\.\d+)?(?:cm|mm|%))?)\s*[：:]?\s*([^】\]]*?)[】\]]'
)


def parse_image_marker(text):
    """
    解析图片标记文本，提取图片名、宽度规格和图题
    
    支持的格式：
      [图片1]          → name="图片1", width_spec="", caption=""
      [图片1,10cm]     → name="图片1", width_spec="10cm", caption=""
      [图片1,50%]      → name="图片1", width_spec="50%", caption=""
      [图片1,宽8cm]    → name="图片1", width_spec="8cm", caption=""
      [图片1:图题]     → name="图片1", width_spec="", caption="图题"
      [图片1,10cm:图题] → name="图片1", width_spec="10cm", caption="图题"
    
    返回：
      (image_name, width_spec, caption) 或 None
    """
    m = RE_IMAGE.match(text.strip())
    if not m:
        return None
    name = m.group(1)
    width_spec = m.group(2).strip()
    caption = m.group(3).strip()
    # 处理"宽8cm"格式，去掉"宽"前缀
    if width_spec.startswith('宽'):
        width_spec = width_spec[1:]
    return name, width_spec, caption


def parse_image_width(width_spec, page_content_width_mm=PAGE_CONTENT_WIDTH_MM):
    """
    解析宽度规格，返回 Cm 对象
    
    参数：
      width_spec: 宽度规格字符串，如 "10cm", "50%", "80mm"
      page_content_width_mm: 页面有效宽度（mm），用于百分比计算
    
    返回：
      Cm 对象，或 None（表示使用默认宽度）
    """
    if not width_spec:
        return None
    
    # 百分比
    m_pct = re.match(r'^(\d+(?:\.\d+)?)%$', width_spec)
    if m_pct:
        pct = float(m_pct.group(1)) / 100.0
        return Mm(page_content_width_mm * pct)
    
    # 厘米
    m_cm = re.match(r'^(\d+(?:\.\d+)?)cm$', width_spec)
    if m_cm:
        return Cm(float(m_cm.group(1)))
    
    # 毫米
    m_mm = re.match(r'^(\d+(?:\.\d+)?)mm$', width_spec)
    if m_mm:
        return Mm(float(m_mm.group(1)))
    
    return None


def parse_text(text):
    """
    解析纯文本，识别各段落类型并结构化返回
    
    返回结构：
    [
        {'type': 'main_title', 'text': '...'},      # 公文主标题
        {'type': 'blank', 'text': ''},               # 空行（标题后空行）
        {'type': 'fawen_hao', 'text': '...'},        # 发文字号
        {'type': 'body', 'text': '...'},             # 正文
        {'type': 'level1', 'text': '...'},           # 一级标题
        {'type': 'level2', 'text': '...'},           # 二级标题
        {'type': 'level3', 'text': '...'},           # 三级标题
        {'type': 'attachment_marker', 'text': '...'}, # 附件标识
        {'type': 'attachment_title', 'text': '...'},  # 附件标题
        {'type': 'table', 'rows': [[...], ...]},     # 表格
        {'type': 'image', 'name': '图片1', 'width_spec': '', 'caption': ''},  # 图片标记
    ]
    """
    lines = text.split('\n')
    elements = []
    i = 0
    main_title_found = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行（但在特定位置保留）
        if not stripped:
            i += 1
            continue

        # --- 公文主标题：第一个非空行，且不超过20字 ---
        if not main_title_found:
            if len(stripped) <= 20:
                elements.append({'type': 'main_title', 'text': stripped})
                # 标题与正文之间空一行
                elements.append({'type': 'blank', 'text': ''})
                main_title_found = True
                i += 1
                continue
            else:
                # 超过20字，仍作为主标题处理（宽松处理）
                elements.append({'type': 'main_title', 'text': stripped})
                elements.append({'type': 'blank', 'text': ''})
                main_title_found = True
                i += 1
                continue

        # --- 图片标记识别（优先于表格识别，避免方括号被误判为表格） ---
        img_parsed = parse_image_marker(stripped)
        if img_parsed:
            name, width_spec, caption = img_parsed
            elements.append({
                'type': 'image',
                'name': name,
                'width_spec': width_spec,
                'caption': caption,
                'text': stripped,  # 保留原始标记文本
            })
            i += 1
            continue

        # --- 表格识别 ---
        if RE_TABLE_ROW.match(stripped):
            table_rows = []
            while i < len(lines):
                row_line = lines[i].strip()
                if not RE_TABLE_ROW.match(row_line):
                    break
                # 跳过分隔行（|---|---|）
                if RE_TABLE_SEPARATOR.match(row_line):
                    i += 1
                    continue
                # 解析单元格
                cells = [c.strip() for c in row_line.split('|') if c.strip() != '']
                table_rows.append(cells)
                i += 1
            if table_rows:
                elements.append({'type': 'table', 'rows': table_rows})
            continue

        # --- 附件识别 ---
        if stripped.startswith('附件') or stripped.startswith('附件：'):
            elements.append({'type': 'blank', 'text': ''})  # 附件前空一行
            elements.append({'type': 'attachment_marker', 'text': stripped})
            i += 1
            continue

        # --- 发文字号 ---
        if RE_FAWEN_HAO.match(stripped):
            elements.append({'type': 'fawen_hao', 'text': stripped})
            i += 1
            continue

        # --- 一级标题 ---
        if RE_LEVEL1.match(stripped):
            elements.append({'type': 'level1', 'text': stripped})
            i += 1
            continue

        # --- 二级标题 ---
        if RE_LEVEL2.match(stripped):
            elements.append({'type': 'level2', 'text': stripped})
            i += 1
            continue

        # --- 三级标题 ---
        if RE_LEVEL3.match(stripped):
            elements.append({'type': 'level3', 'text': stripped})
            i += 1
            continue

        # --- 默认为正文 ---
        elements.append({'type': 'body', 'text': stripped})
        i += 1

    return elements


# ============================================================
# 4. 格式应用模块
# ============================================================

def apply_main_title(doc, text):
    """应用公文主标题格式：方正小标宋简体/黑体，二号，居中，行距28磅"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    # 优先使用方正小标宋简体，若不可用则用黑体替代
    set_run_font(run, FONT_XIAOBIAOSONG, SIZE_ERHAO, bold=False)
    apply_number_font(para)


def apply_blank_line(doc):
    """添加空行（标题与正文之间），字号三号，无缩进"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run('')
    set_run_font(run, FONT_FANGSONG, SIZE_SANHAO)


def apply_fawen_hao(doc, text):
    """应用发文字号格式：仿宋，三号，居中，不加粗"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=False)
    apply_number_font(para)


def apply_body(doc, text, first_indent=True, doc_type="通知"):
    """
    应用正文格式：仿宋_GB2312，三号，首行缩进2字符，行距28磅
    - first_indent: 是否首行缩进，附件正文不缩进
    - doc_type: 公文类型，用于判断问候语是否缩进
    """
    para = doc.add_paragraph()
    indent = Pt(32) if first_indent else None  # 2字符约32磅（三号字2个字符宽度）

    # 判断是否为问候语行（讲话稿/会议纪要中不缩进）
    if first_indent and doc_type in ["讲话稿", "会议纪要"]:
        text_stripped = text.strip()
        if RE_GREETING.match(text_stripped):
            indent = Pt(0)  # 问候语不缩进

    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=indent,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=False)
    apply_number_font(para)


def apply_number_font(paragraph):
    """将段落中的数字和英文字母设置为 Times New Roman 字体，其余保持原字体"""
    runs_to_remove = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        # 用正则分割：数字+英文部分 和 非数字英文部分
        parts = re.split(r'([0-9a-zA-Z]+)', text)
        # 检查是否需要拆分（即是否有混合内容）
        non_empty_parts = [p for p in parts if p]
        has_alnum = any(re.match(r'^[0-9a-zA-Z]+$', p) for p in non_empty_parts)
        has_non_alnum = any(not re.match(r'^[0-9a-zA-Z]+$', p) for p in non_empty_parts)
        if not (has_alnum and has_non_alnum):
            # 无需拆分，但如果全是数字/英文则改字体
            if has_alnum and not has_non_alnum:
                run.font.name = 'Times New Roman'
                # 更新 rFonts 中的 ascii 和 hAnsi 属性
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            continue
        # 需要拆分：为每个 part 创建新 run
        for part in non_empty_parts:
            new_run = paragraph.add_run(part)
            # 复制原 run 的格式
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            if run.font.color.rgb is not None:
                new_run.font.color.rgb = run.font.color.rgb
            # 设置字体
            if re.match(r'^[0-9a-zA-Z]+$', part):
                new_run.font.name = 'Times New Roman'
                # 复制 eastAsia 字体保持原样
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea_font = rFonts.get(qn('w:eastAsia'))
                else:
                    ea_font = None
                nr = new_run._element
                nrPr = nr.get_or_add_rPr()
                nrFonts = nrPr.find(qn('w:rFonts'))
                if nrFonts is None:
                    nrFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    nrPr.insert(0, nrFonts)
                nrFonts.set(qn('w:ascii'), 'Times New Roman')
                nrFonts.set(qn('w:hAnsi'), 'Times New Roman')
                if ea_font:
                    nrFonts.set(qn('w:eastAsia'), ea_font)
            else:
                # 中文部分保持原字体
                new_run.font.name = run.font.name
                # 复制 rFonts 属性
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    nr = new_run._element
                    nrPr = nr.get_or_add_rPr()
                    nrFonts = nrPr.find(qn('w:rFonts'))
                    if nrFonts is None:
                        nrFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                        nrPr.insert(0, nrFonts)
                    for attr_key, attr_val in rFonts.attrib.items():
                        nrFonts.set(attr_key, attr_val)
        # 记录需要删除的原 run
        runs_to_remove.append(run)
    # 删除原来的 run
    for run in runs_to_remove:
        run._element.getparent().remove(run._element)


def apply_level1_title(doc, text):
    """一级标题格式：黑体，三号，首行缩进2字符，行距28磅"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Pt(32),
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    set_run_font(run, FONT_HEITI, SIZE_SANHAO, bold=False)
    apply_number_font(para)


def apply_level2_title(doc, text):
    """二级标题格式：楷体_GB2312/楷体，三号，首行缩进2字符，行距28磅"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Pt(32),
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    set_run_font(run, FONT_KAITI, SIZE_SANHAO, bold=False)
    apply_number_font(para)


def apply_level3_title(doc, text):
    """三级标题格式：仿宋_GB2312加粗，三号，首行缩进2字符，行距28磅"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Pt(32),
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=True)
    apply_number_font(para)


def apply_attachment_marker(doc, text):
    """附件标识格式：仿宋，三号，无首行缩进"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    run = para.add_run(text)
    set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=False)
    apply_number_font(para)


def apply_table(doc, table_data):
    """
    应用表格格式
    - 表格宽度与正文对齐
    - 表头：黑体，三号，居中
    - 表内文字：仿宋，三号，居中
    - 边框：全部实线，黑色，0.5磅
    - 单元格对齐：垂直居中，水平居中
    """
    if not table_data or not table_data[0]:
        return

    num_cols = len(table_data[0])
    num_rows = len(table_data)

    # 计算正文区域宽度（A4宽 - 左右边距）
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    # 添加表格前空行
    para_before = doc.add_paragraph()
    set_paragraph_format(
        para_before,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )

    # 创建表格
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格总宽度
    table_element = table._tbl
    tblPr = table_element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table_element.insert(0, tblPr)

    # 设置表格宽度为正文宽度
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{content_width}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(content_width))
        tblW.set(qn('w:type'), 'dxa')

    # 设置每列等宽
    col_width = int(content_width / num_cols)
    tblGrid = table_element.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
        table_element.insert(1, tblGrid)
    else:
        # 清空已有列定义
        for gc in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(gc)
    for _ in range(num_cols):
        gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{col_width}"/>')
        tblGrid.append(gridCol)

    # 设置表格边框：全部实线，黑色，0.5磅
    set_table_borders(table)

    # 填充表格内容
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            # 清空默认段落
            cell_para = cell.paragraphs[0]
            cell_para.clear()

            # 设置单元格对齐：垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 水平居中
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 设置单元格段落格式
            pf = cell_para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = Pt(28)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

            run = cell_para.add_run(cell_text.strip())

            # 表头（第一行）用黑体，其余用仿宋
            if row_idx == 0:
                set_run_font(run, FONT_HEITI, SIZE_SANHAO, bold=False)
            else:
                set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=False)

    # 表格后空行
    para_after = doc.add_paragraph()
    set_paragraph_format(
        para_after,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )


def set_table_borders(table):
    """
    设置表格边框：全部实线，黑色，0.5磅
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # 移除已有边框设置
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    # 添加边框
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>'''
    # w:sz="4" = 0.5磅 (1/8磅为单位, 4*1/8=0.5磅)
    borders_element = parse_xml(borders_xml)
    tblPr.append(borders_element)


# ============================================================
# 5. 图片处理模块
# ============================================================

def add_image_to_doc(doc, image_data, width=None, caption=None):
    """
    在文档中插入图片
    
    参数：
      doc: Document 对象
      image_data: 图片数据，可以是文件路径(str)或字节流(bytes)
      width: 图片宽度（Cm/Mm 对象），默认为页面有效宽度的80%（约12.5cm）
      caption: 图题文字，如 "图1 网络架构图"
    """
    # 计算默认宽度：页面有效宽度的80%
    if width is None:
        width = Cm(IMAGE_DEFAULT_WIDTH_CM)
    
    # 图片前空0.5行（约14磅）
    para_before = doc.add_paragraph()
    set_paragraph_format(
        para_before,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )
    
    # 添加图片段落
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    
    # 处理不同类型的图片数据
    if isinstance(image_data, bytes):
        image_stream = io.BytesIO(image_data)
        run.add_picture(image_stream, width=width)
    elif isinstance(image_data, str):
        # 文件路径
        run.add_picture(image_data, width=width)
    else:
        # 假设是文件类对象（如 Streamlit UploadedFile）
        image_stream = io.BytesIO(image_data.read())
        run.add_picture(image_stream, width=width)
    
    # 图片段落的段落格式
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    
    # 添加图题
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.font.name = FONT_FANGSONG_ALT
        # 设置东亚字体
        r = cap_run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_FANGSONG_ALT}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), FONT_FANGSONG_ALT)
        cap_run.font.size = SIZE_XIAOSI  # 小四号（12磅）
        cap_para.paragraph_format.space_before = Pt(7)   # 约0.5行
        cap_para.paragraph_format.space_after = Pt(7)
        cap_para.paragraph_format.line_spacing = LINE_SPACING
        cap_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    
    # 图片后空0.5行（约14磅）
    para_after = doc.add_paragraph()
    set_paragraph_format(
        para_after,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        line_spacing=LINE_SPACING,
        space_before=Pt(0),
        space_after=Pt(0)
    )


def process_image_element(doc, elem, images):
    """
    处理图片类型的元素
    
    参数：
      doc: Document 对象
      elem: 解析后的图片元素 {'type': 'image', 'name': '图片1', ...}
      images: 图片字典，key为标记名如"图片1"，value为图片数据
    """
    name = elem.get('name', '')
    width_spec = elem.get('width_spec', '')
    caption = elem.get('caption', '')
    
    # 查找图片数据
    image_data = None
    if images and name in images:
        image_data = images[name]
    
    if image_data is None:
        # 没有对应图片数据，将标记作为普通正文输出
        apply_body(doc, elem.get('text', ''), first_indent=True)
        return
    
    # 解析宽度
    width = parse_image_width(width_spec)
    
    # 自动生成图题编号
    # 从图片名中提取数字，如 "图片1" → "图1"
    num_match = re.search(r'\d+', name)
    img_num = num_match.group(0) if num_match else ''
    
    # 如果有图题文字，添加编号前缀
    if caption and img_num:
        caption = f"图{img_num} {caption}"
    elif caption:
        caption = caption
    
    # 插入图片
    add_image_to_doc(doc, image_data, width=width, caption=caption if caption else None)


# ============================================================
# 6. 文件输出模块
# ============================================================

def format_gongwen(text, output_path=None, images=None, doc_type="通知"):
    """
    主格式化函数：接收纯文本，输出规范格式的 .docx 文件
    
    参数：
      text: 纯文本字符串
      output_path: 输出文件路径，默认为 "公文_时间戳.docx"
      images: 图片字典，key为标记名如"图片1"，value为图片数据(bytes/文件路径)
      doc_type: 公文类型，如"通知"、"讲话稿"等
    
    返回：
      生成的文件路径
    """
    # 创建文档
    doc = Document()

    # 1. 页面设置
    setup_page(doc)

    # 2. 解析文本
    elements = parse_text(text)

    # 3. 应用格式
    for elem in elements:
        etype = elem['type']

        if etype == 'main_title':
            apply_main_title(doc, elem['text'])

        elif etype == 'blank':
            apply_blank_line(doc)

        elif etype == 'fawen_hao':
            apply_fawen_hao(doc, elem['text'])

        elif etype == 'body':
            apply_body(doc, elem['text'], first_indent=True, doc_type=doc_type)

        elif etype == 'level1':
            apply_level1_title(doc, elem['text'])

        elif etype == 'level2':
            apply_level2_title(doc, elem['text'])

        elif etype == 'level3':
            apply_level3_title(doc, elem['text'])

        elif etype == 'attachment_marker':
            apply_attachment_marker(doc, elem['text'])

        elif etype == 'table':
            apply_table(doc, elem['rows'])

        elif etype == 'image':
            process_image_element(doc, elem, images)

    # 4. 添加页码
    add_page_number(doc)

    # 5. 输出文件
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f'公文_{timestamp}.docx'

    # 确保目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    doc.save(output_path)
    return output_path


# ============================================================
# 入口函数
# ============================================================

def main(text, output_path=None, images=None, doc_type="通知"):
    """
    入口函数
    
    参数：
      text: 纯文本字符串（公文内容）
      output_path: 输出文件路径（可选）
      images: 图片字典，key为标记名如"图片1"，value为图片数据(bytes或文件路径)（可选）
      doc_type: 公文类型，默认"通知"
    
    返回：
      生成的文件路径
    """
    return format_gongwen(text, output_path, images=images, doc_type=doc_type)


if __name__ == '__main__':
    import sys

    if len(sys.argv) >= 2:
        # 从命令行参数获取文本（或文件路径）
        arg = sys.argv[1]
        if os.path.isfile(arg):
            with open(arg, 'r', encoding='utf-8') as f:
                input_text = f.read()
        else:
            input_text = arg

        out = sys.argv[2] if len(sys.argv) >= 3 else None
        result = main(input_text, out)
        print(f'公文已生成：{result}')
    else:
        print('用法：python gongwen_formatter.py <文本或文件路径> [输出路径]')
        print('示例：python gongwen_formatter.py "公文内容.txt" "输出.docx"')
